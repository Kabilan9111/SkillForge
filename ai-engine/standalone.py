"""
SkillForge AI Engine — Self-Contained Local Pipeline
====================================================
Runs WITHOUT any external dependencies:
  ✓ No OpenAI / Anthropic API keys required
  ✓ No PostgreSQL required
  ✓ No PyTorch / spaCy / FAISS required
  ✓ No Neo4j required

Dependencies (all lightweight, already in requirements.txt):
  fastapi, uvicorn, PyMuPDF, python-docx, python-multipart, python-jose

Start:
  python -m uvicorn standalone:app --host 0.0.0.0 --port 8001 --reload
"""

from __future__ import annotations

import io
import json
import logging
import os
import re
import time
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Optional

from fastapi import FastAPI, File, Form, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s — %(message)s",
)
log = logging.getLogger("skillforge.local")

# ══════════════════════════════════════════════════════════════════════
# LAYER 1 — DOCUMENT TEXT EXTRACTION
# ══════════════════════════════════════════════════════════════════════

def extract_text_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF using pdfminer.six — pure Python, no compilation."""
    try:
        from pdfminer.high_level import extract_text as pdfminer_extract
        return pdfminer_extract(io.BytesIO(file_bytes)) or ""
    except ImportError:
        raise RuntimeError("pdfminer.six not installed. Run: pip install pdfminer.six")
    except Exception as e:
        raise RuntimeError(f"PDF extraction failed: {e}")


def extract_text_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX using python-docx."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)
        return "\n".join(paragraphs)
    except ImportError:
        raise RuntimeError("python-docx not installed. Run: pip install python-docx")
    except Exception as e:
        raise RuntimeError(f"DOCX extraction failed: {e}")


def extract_text(file_bytes: bytes, filename: str) -> str:
    """Route to correct extractor based on file extension."""
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return extract_text_pdf(file_bytes)
    elif ext in (".docx", ".doc"):
        return extract_text_docx(file_bytes)
    elif ext == ".txt":
        return file_bytes.decode("utf-8", errors="replace")
    else:
        # Try PDF first, then DOCX, then raw UTF-8
        try:
            return extract_text_pdf(file_bytes)
        except Exception:
            try:
                return extract_text_docx(file_bytes)
            except Exception:
                return file_bytes.decode("utf-8", errors="replace")


# ══════════════════════════════════════════════════════════════════════
# LAYER 2 — SKILL EXTRACTION (Ontology-Based Matching)
# ══════════════════════════════════════════════════════════════════════

# ── Master Skill Ontology (500+ entries, categorised) ──────────────────
SKILL_ONTOLOGY: dict[str, list[str]] = {
    "Programming Languages": [
        "Python","JavaScript","TypeScript","Java","C++","C#","C","Go","Golang",
        "Rust","Ruby","PHP","Swift","Kotlin","Scala","R","MATLAB","Perl","Haskell",
        "Elixir","Erlang","Dart","Lua","Groovy","Clojure","F#","OCaml",
        "Assembly","COBOL","Fortran","Julia","Shell","Bash","PowerShell",
        "SQL","PL/SQL","T-SQL","PLSQL","HCL","YAML","Solidity","Move",
    ],
    "Frontend Frameworks": [
        "React","ReactJS","React.js","Next.js","NextJS","Vue","Vue.js","Angular",
        "AngularJS","Svelte","SvelteKit","Ember","Nuxt","Nuxt.js","Gatsby",
        "Remix","Astro","Qwik","Lit","Alpine.js","jQuery","Bootstrap",
        "Tailwind CSS","TailwindCSS","Material UI","MUI","Chakra UI","Ant Design",
        "Storybook","WebPack","Vite","Parcel","Rollup","Babel","ESLint",
        "Redux","Zustand","MobX","Recoil","Jotai","React Query","SWR",
        "React Native","Expo","Ionic","Capacitor","Electron",
    ],
    "Backend Frameworks": [
        "Node.js","NodeJS","Express","Express.js","NestJS","Fastify","Koa",
        "Hapi","Django","Flask","FastAPI","Tornado","Aiohttp","Sanic",
        "Spring","Spring Boot","Spring MVC","Spring Security","Spring Data",
        "Hibernate","JPA","Jakarta EE","Quarkus","Micronaut","Vert.x",
        "Ruby on Rails","Rails","Sinatra","Hanami",
        "Laravel","Symfony","CodeIgniter","Lumen",
        "ASP.NET","ASP.NET Core",".NET","gRPC","GraphQL","REST","REST API",
        "WebSockets","Socket.io","Celery","FastAPI","Pydantic","SQLAlchemy",
        "Tortoise ORM","Prisma","TypeORM","Sequelize","Mongoose","Drizzle",
    ],
    "Databases": [
        "PostgreSQL","Postgres","MySQL","MariaDB","SQLite","SQL Server","MSSQL",
        "Oracle","DB2","MongoDB","Mongoose","Cassandra","DynamoDB","Couchbase",
        "CouchDB","RethinkDB","ArangoDB","Neo4j","TigerGraph","Amazon Neptune",
        "Redis","Memcached","KeyDB","Dragonfly",
        "Elasticsearch","OpenSearch","Solr","Algolia","Typesense",
        "ClickHouse","BigQuery","Snowflake","Redshift","Databricks","Delta Lake",
        "Supabase","PlanetScale","Neon","CockroachDB","TiDB","YugabyteDB",
        "InfluxDB","TimescaleDB","Prometheus","VictoriaMetrics",
        "Pinecone","Weaviate","Chroma","Qdrant","Milvus","FAISS",
    ],
    "Cloud & Infrastructure": [
        "AWS","Amazon Web Services","GCP","Google Cloud","Azure","Microsoft Azure",
        "DigitalOcean","Linode","Hetzner","OVH","Cloudflare",
        "EC2","S3","Lambda","ECS","EKS","RDS","SQS","SNS","SES","API Gateway",
        "CloudFront","Route 53","IAM","VPC","ELB","ALB","NLB",
        "GKE","GAE","Cloud Run","Cloud Functions","Cloud Storage","BigQuery",
        "Azure AKS","Azure Functions","Azure DevOps","Azure Blob","Azure SQL",
        "Terraform","Terraform Cloud","Pulumi","CDK","AWS CDK","Ansible","Chef",
        "Puppet","SaltStack","Vagrant","Packer","Nomad","Consul",
    ],
    "DevOps & CI/CD": [
        "Docker","Docker Compose","Kubernetes","K8s","Helm","Kustomize",
        "Istio","Linkerd","Envoy","Nginx","Apache","Caddy","Traefik",
        "Jenkins","GitLab CI","GitHub Actions","CircleCI","Travis CI",
        "TeamCity","Bamboo","Azure DevOps","ArgoCD","Flux","Spinnaker",
        "Datadog","Grafana","Prometheus","Jaeger","Zipkin","OpenTelemetry",
        "ELK","Elasticsearch","Logstash","Kibana","Splunk","New Relic",
        "PagerDuty","OpsGenie","VictorOps","SRE","GitOps","DevSecOps","DevOps",
        "CI/CD","Continuous Integration","Continuous Deployment","Infrastructure as Code",
        "Git","GitHub","GitLab","Bitbucket","SVN","Mercurial",
    ],
    "AI & Machine Learning": [
        "Machine Learning","ML","Deep Learning","Neural Networks","NLP",
        "Computer Vision","Reinforcement Learning","Federated Learning",
        "PyTorch","TensorFlow","Keras","JAX","Flax","MXNet","Caffe",
        "Scikit-learn","XGBoost","LightGBM","CatBoost","H2O","AutoML",
        "HuggingFace","Transformers","BERT","GPT","LLaMA","Mistral",
        "LangChain","LlamaIndex","Semantic Kernel","AutoGen","CrewAI",
        "OpenAI","Anthropic","Cohere","Vertex AI","SageMaker","Azure ML",
        "MLflow","Weights & Biases","W&B","DVC","ZenML","Kubeflow","Metaflow",
        "ONNX","TensorRT","Triton","vLLM","TorchServe","BentoML","Ray",
        "Pandas","NumPy","SciPy","Matplotlib","Seaborn","Plotly","Bokeh",
        "Spark MLlib","Dask","Rapids","CuDF","Intel oneAPI",
        "Data Science","Data Analysis","Statistics","A/B Testing",
        "Feature Engineering","Hyperparameter Tuning","Model Deployment",
    ],
    "Data Engineering": [
        "Apache Spark","Spark","PySpark","Delta Lake","Apache Kafka","Kafka",
        "Apache Flink","Flink","Apache Beam","Apache Airflow","Airflow",
        "Prefect","Dagster","Luigi","dbt","dbt Cloud","Lightdash",
        "Snowflake","BigQuery","Redshift","Databricks","Delta Lake",
        "Apache Hive","Presto","Trino","Athena","Apache Iceberg","Apache Hudi",
        "Fivetran","Airbyte","Stitch","Talend","Informatica","SSIS",
        "ETL","ELT","Data Modeling","Data Warehouse","Data Lake","Lakehouse",
        "Data Pipeline","Stream Processing","Batch Processing","CDC",
    ],
    "Security": [
        "Cybersecurity","Security","OWASP","Penetration Testing","Pen Testing",
        "Ethical Hacking","Bug Bounty","SAST","DAST","IAST","SCA",
        "Vulnerability Assessment","Threat Modeling","Security Auditing",
        "OAuth","OAuth2","JWT","SAML","OIDC","SSO","MFA","2FA",
        "Encryption","PKI","TLS","SSL","HTTPS","Zero Trust","RBAC","ABAC",
        "Snyk","SonarQube","Checkmarx","Veracode","Bandit","Semgrep",
        "Vault","AWS Secrets Manager","KMS","HSM","Key Management",
        "SIEM","SOC","IDS","IPS","WAF","Firewall","VPN","Zero Knowledge",
    ],
    "Methodologies & Practices": [
        "Agile","Scrum","Kanban","SAFe","Lean","XP","TDD","BDD","DDD",
        "Microservices","Monolith","SOA","Event-Driven Architecture","CQRS",
        "Event Sourcing","SAGA","Hexagonal Architecture","Clean Architecture",
        "System Design","Distributed Systems","Scalability","High Availability",
        "API Design","OpenAPI","Swagger","Postman","Insomnia",
        "Unit Testing","Integration Testing","E2E Testing","Load Testing",
        "Jest","Vitest","Pytest","JUnit","Mocha","Chai","Cypress","Playwright",
        "Selenium","TestNG","Mockito","Enzyme","React Testing Library",
        "Code Review","Pair Programming","Technical Writing","Documentation",
        "SOLID","Design Patterns","Refactoring","Performance Optimization",
    ],
    "Mobile Development": [
        "Android","iOS","React Native","Flutter","Dart","Swift","Kotlin",
        "Objective-C","Xamarin","Ionic","Cordova","Capacitor","Expo",
        "SwiftUI","Jetpack Compose","RxSwift","RxKotlin","Room","Retrofit",
        "Firebase","Push Notifications","App Store","Play Store",
    ],
    "Networking & Protocols": [
        "TCP/IP","HTTP","HTTPS","HTTP/2","HTTP/3","QUIC","gRPC","WebSocket",
        "MQTT","AMQP","WebRTC","DNS","CDN","Load Balancing","Reverse Proxy",
        "REST","SOAP","GraphQL","JSON","XML","Protocol Buffers","Protobuf",
        "OpenAPI","Swagger","AsyncAPI","Networking","OSI Model","BGP","OSPF",
    ],
}

# Build a flat lookup: lowercase_term → (canonical_name, category)
_SKILL_LOOKUP: dict[str, tuple[str, str]] = {}
for _cat, _terms in SKILL_ONTOLOGY.items():
    for _term in _terms:
        _SKILL_LOOKUP[_term.lower()] = (_term, _cat)

# Build sorted list by length desc so longer matches win (e.g. "Next.js" before "js")
_SKILL_PATTERNS = sorted(_SKILL_LOOKUP.keys(), key=len, reverse=True)


def extract_skills(text: str) -> dict[str, Any]:
    """
    Fast regex scan of resume text against the skill ontology.
    Returns categorised skills + flat list, sorted by frequency.
    No LLM, no API, no internet required.
    """
    text_lower = text.lower()
    matched: dict[str, str] = {}  # canonical → category

    for term_lower in _SKILL_PATTERNS:
        canonical, category = _SKILL_LOOKUP[term_lower]
        # Word-boundary aware match (handles C++, C#, .NET specially)
        if re.search(
            r"(?<![a-z])" + re.escape(term_lower) + r"(?![a-z])",
            text_lower,
        ):
            matched[canonical] = category

    # Group by category
    by_category: dict[str, list[str]] = {}
    for skill, cat in matched.items():
        by_category.setdefault(cat, []).append(skill)

    return {
        "skills": sorted(matched.keys()),
        "by_category": by_category,
        "total": len(matched),
    }


def extract_years_experience(text: str) -> float:
    """Extract years of experience from resume text."""
    patterns = [
        r"(\d+)\+?\s*years?\s+(?:of\s+)?(?:professional\s+)?experience",
        r"experience[:\s]+(\d+)\+?\s*years?",
        r"(\d+)\+?\s*years?\s+in\s+(?:the\s+)?(?:industry|field|software|tech)",
        r"over\s+(\d+)\s+years?",
    ]
    for pat in patterns:
        m = re.search(pat, text, re.I)
        if m:
            return float(m.group(1))
    return 0.0


def extract_contact(text: str) -> dict[str, str]:
    """Extract contact information from resume text."""
    contact: dict[str, str] = {}

    email_m = re.search(r"[\w.+-]+@[\w-]+\.[\w.]+", text)
    if email_m:
        contact["email"] = email_m.group(0)

    phone_m = re.search(r"(?:\+?\d[\d\s\-().]{7,14}\d)", text)
    if phone_m:
        contact["phone"] = phone_m.group(0).strip()

    linkedin_m = re.search(r"linkedin\.com/in/([\w-]+)", text, re.I)
    if linkedin_m:
        contact["linkedin"] = f"linkedin.com/in/{linkedin_m.group(1)}"

    github_m = re.search(r"github\.com/([\w-]+)", text, re.I)
    if github_m:
        contact["github"] = f"github.com/{github_m.group(1)}"

    return contact


# ══════════════════════════════════════════════════════════════════════
# LAYER 3 — STRUCTURED INPUT BUILDING
# ══════════════════════════════════════════════════════════════════════

def build_structured_context(
    user_skills: list[str],
    target_role: str,
    years_exp: float,
) -> str:
    """Build the structured context block (would be sent to LLM if keys were set)."""
    skills_block = "\n".join(f"  • {s}" for s in sorted(user_skills)) or "  (none detected)"
    return (
        f"User Skills:\n{skills_block}\n\n"
        f"Target Role: {target_role}\n\n"
        f"Experience: {years_exp:.0f} year{'s' if years_exp != 1 else ''}\n"
    )


# ══════════════════════════════════════════════════════════════════════
# LAYER 4 — MARKET INTELLIGENCE (Built-in Role Database)
# ══════════════════════════════════════════════════════════════════════

ROLE_DB: dict[str, dict] = {
    "software engineer": {
        "display": "Software Engineer",
        "critical":   ["Data Structures","Algorithms","SQL","Git","Testing","REST API","System Design"],
        "moderate":   ["Docker","CI/CD","Microservices","PostgreSQL","Redis"],
        "nice":       ["Kubernetes","GraphQL","TypeScript","AWS"],
        "salary_range": "$90k – $140k",
    },
    "backend engineer": {
        "display": "Backend Engineer",
        "critical":   ["Python","PostgreSQL","Redis","REST API","Docker","Microservices","Git"],
        "moderate":   ["Kafka","gRPC","Kubernetes","System Design","Terraform","CI/CD"],
        "nice":       ["Elasticsearch","MongoDB","AWS","gRPC"],
        "salary_range": "$100k – $160k",
    },
    "frontend engineer": {
        "display": "Frontend Engineer",
        "critical":   ["React","TypeScript","HTML/CSS","JavaScript","Git","Testing"],
        "moderate":   ["Next.js","State Management","Performance Optimization","Jest"],
        "nice":       ["GraphQL","React Native","Accessibility","WebPack"],
        "salary_range": "$90k – $150k",
    },
    "fullstack engineer": {
        "display": "Fullstack Engineer",
        "critical":   ["React","Node.js","TypeScript","PostgreSQL","REST API","Git"],
        "moderate":   ["Docker","Next.js","Redis","Testing","CI/CD","AWS"],
        "nice":       ["GraphQL","System Design","Kubernetes","Terraform"],
        "salary_range": "$100k – $160k",
    },
    "devops engineer": {
        "display": "DevOps Engineer",
        "critical":   ["Docker","Kubernetes","Terraform","CI/CD","AWS","Linux","Git"],
        "moderate":   ["Ansible","Prometheus","Grafana","Helm","Bash","Python"],
        "nice":       ["GitOps","Service Mesh","Vault","Cost Optimization"],
        "salary_range": "$110k – $170k",
    },
    "cloud engineer": {
        "display": "Cloud Engineer",
        "critical":   ["AWS","Terraform","Kubernetes","Networking","IAM","Linux"],
        "moderate":   ["GCP","Azure","Serverless","Cloud Architecture","Security","Python"],
        "nice":       ["FinOps","Multi-cloud","Compliance","Pulumi"],
        "salary_range": "$110k – $170k",
    },
    "data engineer": {
        "display": "Data Engineer",
        "critical":   ["Python","SQL","Apache Spark","Apache Airflow","Data Modeling","Git"],
        "moderate":   ["Kafka","dbt","Snowflake","BigQuery","ETL","Docker"],
        "nice":       ["Scala","Flink","Delta Lake","Data Governance","Terraform"],
        "salary_range": "$100k – $160k",
    },
    "data scientist": {
        "display": "Data Scientist",
        "critical":   ["Python","Machine Learning","Statistics","Scikit-learn","SQL","Pandas"],
        "moderate":   ["XGBoost","Deep Learning","Feature Engineering","A/B Testing","NumPy"],
        "nice":       ["Apache Spark","MLflow","PyTorch","TensorFlow","Docker"],
        "salary_range": "$100k – $160k",
    },
    "machine learning engineer": {
        "display": "Machine Learning Engineer",
        "critical":   ["Python","PyTorch","TensorFlow","Scikit-learn","SQL","Docker","Git"],
        "moderate":   ["MLflow","Kubernetes","FastAPI","REST API","NumPy","Pandas","AWS"],
        "nice":       ["LLM","LangChain","HuggingFace","Model Serving","CUDA","Ray"],
        "salary_range": "$120k – $180k",
    },
    "ai engineer": {
        "display": "AI Engineer",
        "critical":   ["Python","PyTorch","LangChain","FastAPI","Docker","Git","LLM"],
        "moderate":   ["HuggingFace","MLflow","REST API","Kubernetes","AWS","Vector Databases"],
        "nice":       ["CUDA","Triton","Model Serving","Fine-tuning","RAG","vLLM"],
        "salary_range": "$130k – $200k",
    },
    "senior software engineer": {
        "display": "Senior Software Engineer",
        "critical":   ["System Design","Distributed Systems","PostgreSQL","AWS","Git","Docker","CI/CD"],
        "moderate":   ["Kafka","Redis","Microservices","Performance Tuning","Kubernetes"],
        "nice":       ["Cloud Architecture","Cost Optimization","Technical Leadership","SRE"],
        "salary_range": "$140k – $220k",
    },
    "site reliability engineer": {
        "display": "Site Reliability Engineer",
        "critical":   ["Linux","Kubernetes","Docker","Python","Terraform","Prometheus","CI/CD"],
        "moderate":   ["Grafana","Alerting","Incident Management","AWS","SLO","SLA","Bash"],
        "nice":       ["Chaos Engineering","eBPF","Distributed Tracing","Istio"],
        "salary_range": "$130k – $200k",
    },
    "mobile developer": {
        "display": "Mobile Developer",
        "critical":   ["React Native","JavaScript","TypeScript","iOS","Android","Git"],
        "moderate":   ["Expo","Redux","REST API","Testing","Firebase","App Store"],
        "nice":       ["Swift","Kotlin","Flutter","CI/CD","Push Notifications"],
        "salary_range": "$90k – $150k",
    },
    "ios developer": {
        "display": "iOS Developer",
        "critical":   ["Swift","SwiftUI","UIKit","Xcode","iOS","Git","REST API"],
        "moderate":   ["Objective-C","Core Data","Testing","CI/CD","App Store","Firebase"],
        "nice":       ["Combine","RxSwift","Metal","ARKit","CoreML"],
        "salary_range": "$100k – $160k",
    },
    "android developer": {
        "display": "Android Developer",
        "critical":   ["Kotlin","Android","Jetpack Compose","Android SDK","Git","REST API"],
        "moderate":   ["Java","Room","Retrofit","Testing","CI/CD","Firebase"],
        "nice":       ["RxKotlin","Coroutines","Dagger","Hilt","MVVM"],
        "salary_range": "$100k – $160k",
    },
    "security engineer": {
        "display": "Security Engineer",
        "critical":   ["Cybersecurity","Python","Linux","OWASP","Penetration Testing","Git"],
        "moderate":   ["AWS Security","Vault","TLS","OAuth","SAST","DAST","Networking"],
        "nice":       ["Bug Bounty","Threat Modeling","Zero Trust","SIEM","DevSecOps"],
        "salary_range": "$120k – $180k",
    },
    "product manager": {
        "display": "Product Manager",
        "critical":   ["Product Strategy","Agile","Scrum","User Research","Roadmapping","SQL"],
        "moderate":   ["Data Analysis","A/B Testing","Stakeholder Management","Jira"],
        "nice":       ["Technical Knowledge","Python","System Design","Figma"],
        "salary_range": "$110k – $170k",
    },
}


def match_role(target_role: str) -> tuple[str, dict | None]:
    """Match target role to our database (fuzzy)."""
    key = target_role.strip().lower()
    if key in ROLE_DB:
        return key, ROLE_DB[key]

    # Fuzzy: check if any DB key is a substring or vice versa
    for db_key, data in ROLE_DB.items():
        if db_key in key or key in db_key:
            return db_key, data
        # Word overlap
        db_words = set(db_key.split())
        req_words = set(key.split())
        if len(db_words & req_words) >= 2:
            return db_key, data

    return key, None


def compute_skill_gap(
    user_skills: list[str],
    role_data: dict,
) -> dict[str, Any]:
    """Compare user skills against role requirements."""
    user_lower = {s.lower() for s in user_skills}

    def _matched(req_list: list[str]) -> tuple[list[str], list[str]]:
        have, missing = [], []
        for skill in req_list:
            if skill.lower() in user_lower:
                have.append(skill)
            else:
                # Partial match: e.g. user has "PostgreSQL" and req is "SQL"
                partial = any(
                    skill.lower() in u or u in skill.lower()
                    for u in user_lower
                )
                if partial:
                    have.append(skill)
                else:
                    missing.append(skill)
        return have, missing

    crit_have, crit_missing = _matched(role_data["critical"])
    mod_have,  mod_missing  = _matched(role_data["moderate"])
    nice_have, nice_missing = _matched(role_data["nice"])

    all_missing = crit_missing + mod_missing + nice_missing
    all_required = role_data["critical"] + role_data["moderate"] + role_data["nice"]
    all_covered = len(crit_have) + len(mod_have) + len(nice_have)

    # Weighted role alignment: critical=3pts, moderate=2pts, nice=1pt
    max_score = len(role_data["critical"]) * 3 + len(role_data["moderate"]) * 2 + len(role_data["nice"]) * 1
    earned    = len(crit_have) * 3 + len(mod_have) * 2 + len(nice_have) * 1
    alignment = round((earned / max_score * 100) if max_score > 0 else 0)

    return {
        "critical_missing":  crit_missing,
        "moderate_missing":  mod_missing,
        "nice_to_have_missing": nice_missing,
        "missing_skills":    all_missing,
        "covered_count":     all_covered,
        "total_required":    len(all_required),
        "role_alignment_score": alignment,
    }


# ══════════════════════════════════════════════════════════════════════
# LAYER 5 — CONFIDENCE SCORING
# ══════════════════════════════════════════════════════════════════════

def compute_confidence(
    text_length: int,
    skills_found: int,
    extraction_method: str,  # "pdf" | "docx" | "txt"
    years_exp: float,
    has_contact: bool,
    role_matched: bool,
    alignment_score: int,
) -> dict[str, Any]:
    """
    Rule-based confidence score 0–100.
    Based on: text quality, skill density, role match, contact presence.
    """
    score = 0
    breakdown: dict[str, int] = {}

    # Resume text completeness (0-25 pts)
    if text_length >= 2000:
        t = 25
    elif text_length >= 1000:
        t = 18
    elif text_length >= 500:
        t = 12
    elif text_length >= 200:
        t = 6
    else:
        t = 0
    score += t
    breakdown["text_completeness"] = t

    # Skill detection quality (0-35 pts)
    if skills_found >= 20:
        s = 35
    elif skills_found >= 12:
        s = 28
    elif skills_found >= 7:
        s = 20
    elif skills_found >= 4:
        s = 12
    elif skills_found >= 1:
        s = 5
    else:
        s = 0
    score += s
    breakdown["skill_detection"] = s

    # Role match (0-20 pts)
    r = 20 if role_matched else 5
    score += r
    breakdown["role_match"] = r

    # Contact info present (0-10 pts)
    c = 10 if has_contact else 0
    score += c
    breakdown["contact_info"] = c

    # Experience mentioned (0-10 pts)
    e = 10 if years_exp > 0 else 3
    score += e
    breakdown["experience_mentioned"] = e

    score = min(100, score)

    warnings: list[str] = []
    if score < 40:
        warnings.append("Low confidence analysis due to insufficient resume data.")
    if skills_found < 3:
        warnings.append("Very few technical skills detected — ensure resume lists technologies clearly.")
    if text_length < 200:
        warnings.append("Resume text is very short — the file may be image-based or empty.")
    if not role_matched:
        warnings.append("Target role was not found in our database — using generic analysis.")

    if score >= 80:
        grade = "A"
    elif score >= 60:
        grade = "B"
    elif score >= 40:
        grade = "C"
    else:
        grade = "F"

    return {
        "confidence_score": score,
        "grade": grade,
        "breakdown": breakdown,
        "warnings": warnings,
    }


# ══════════════════════════════════════════════════════════════════════
# LAYER 6 — ARBITRATION & FINAL REPORT
# ══════════════════════════════════════════════════════════════════════

def generate_recommendations(
    user_skills: list[str],
    missing_skills: list[str],
    critical_missing: list[str],
    target_role: str,
    alignment_score: int,
    years_exp: float,
    confidence: int,
) -> list[str]:
    """Generate actionable career recommendations."""
    recs: list[str] = []

    # Critical gap recommendations
    if critical_missing:
        top = critical_missing[:3]
        recs.append(
            f"Priority: Learn {', '.join(top)} — these are critical requirements "
            f"for {target_role} roles and will have the highest impact on your hiring chances."
        )

    # Alignment score recommendations
    if alignment_score >= 80:
        recs.append(
            f"You have strong alignment ({alignment_score}%) with {target_role}. "
            f"Focus on deepening expertise in your existing skills and targeting senior-level positions."
        )
    elif alignment_score >= 60:
        recs.append(
            f"Your profile ({alignment_score}% match) is a solid foundation for {target_role}. "
            f"Building on {', '.join((critical_missing[:2] or ['the missing skills']))[:60]} "
            f"will make you competitive."
        )
    elif alignment_score >= 40:
        recs.append(
            f"You're partially aligned ({alignment_score}%) with {target_role}. "
            f"A 2–3 month focused upskilling effort on the critical gaps should make you job-ready."
        )
    else:
        recs.append(
            f"You have a significant skill gap ({alignment_score}% match) for {target_role}. "
            f"Consider an intermediate role or a structured learning plan of 6–12 months."
        )

    # Experience-based
    if years_exp < 1:
        recs.append(
            "Build a strong portfolio with 2–3 projects that demonstrate end-to-end skills — "
            "this compensates for limited work experience when applying."
        )
    elif years_exp < 3:
        recs.append(
            "With your experience level, focus on breadth: touch all layers of the stack, "
            "contribute to open-source, and get AWS/GCP certification."
        )
    else:
        recs.append(
            "With significant experience, emphasise system design, architectural decisions, "
            "and leadership signals in your resume to target senior roles."
        )

    # Skill count
    if len(user_skills) < 5:
        recs.append(
            "Your resume lists very few skills. Expand it with all technologies you have "
            "worked with, even briefly — ATS systems filter heavily on keywords."
        )

    # Certification
    cert_skills = {"AWS", "GCP", "Azure", "Kubernetes", "Docker", "Security"}
    relevant = cert_skills & {s for s in user_skills}
    if relevant:
        recs.append(
            f"Given your {', '.join(list(relevant)[:2])} background, a professional certification "
            f"(e.g. AWS Solutions Architect, CKA) would strengthen your resume significantly."
        )

    return recs[:6]  # Cap at 6 recommendations


def build_final_report(
    user_skills: list[str],
    skills_by_category: dict[str, list[str]],
    gap_result: dict,
    confidence_result: dict,
    target_role: str,
    role_data: dict | None,
    years_exp: float,
    contact: dict,
    text_length: int,
    filename: str,
) -> dict[str, Any]:
    """Combine all layers into the final structured API response."""

    alignment = gap_result.get("role_alignment_score", 0)
    missing = gap_result.get("missing_skills", [])
    critical_missing = gap_result.get("critical_missing", [])

    recs = generate_recommendations(
        user_skills, missing, critical_missing,
        target_role, alignment, years_exp,
        confidence_result["confidence_score"],
    )

    return {
        # Core outputs (matches required spec)
        "user_skills":             user_skills,
        "missing_skills":          missing,
        "role_alignment_score":    alignment,
        "confidence_score":        confidence_result["confidence_score"],
        "career_recommendations":  recs,

        # Extended outputs
        "target_role":             target_role,
        "confirmed_role":          role_data["display"] if role_data else target_role,
        "salary_range":            role_data["salary_range"] if role_data else "Not available",
        "years_experience":        years_exp,

        "skills_by_category":      skills_by_category,
        "total_skills_detected":   len(user_skills),
        "critical_gaps":           critical_missing,
        "moderate_gaps":           gap_result.get("moderate_missing", []),
        "nice_to_have_gaps":       gap_result.get("nice_to_have_missing", []),
        "skills_covered":          gap_result.get("covered_count", 0),
        "total_role_requirements": gap_result.get("total_required", 0),

        "confidence_grade":        confidence_result["grade"],
        "confidence_breakdown":    confidence_result["breakdown"],
        "warnings":                confidence_result["warnings"],
        "contact":                 contact,

        "pipeline_version":        "local-v1.0",
        "analysis_timestamp":      datetime.now(timezone.utc).isoformat(),
        "resume_filename":         filename,
        "text_length":             text_length,
    }


# ══════════════════════════════════════════════════════════════════════
# FASTAPI APP
# ══════════════════════════════════════════════════════════════════════

app = FastAPI(
    title="SkillForge AI Engine (Local)",
    version="1.0.0",
    description=(
        "Self-contained 6-layer AI pipeline. "
        "No API keys, no PostgreSQL, no ML frameworks required."
    ),
    docs_url="/docs",
    redoc_url="/redoc",
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.get("/health", tags=["System"])
async def health():
    return {
        "status": "healthy",
        "engine": "SkillForge Local Pipeline v1.0",
        "layers": ["text_extraction", "skill_ontology", "structured_context",
                   "market_intelligence", "confidence_scoring", "arbitration"],
        "timestamp": datetime.now(timezone.utc).isoformat(),
    }


@app.get("/api/health", tags=["System"])
async def api_health():
    return await health()


@app.post(
    "/api/pipeline/analyze",
    tags=["Pipeline"],
    summary="Full 6-layer AI pipeline — resume → skill gap report",
)
async def analyze_pipeline(
    resume: UploadFile = File(..., description="Resume PDF, DOCX, or TXT"),
    target_role: str = Form(..., description="Target job title"),
    target_salary: Optional[float] = Form(None),
    location: str = Form("United States"),
    years_of_experience: Optional[float] = Form(None),
    github_username: Optional[str] = Form(None),
):
    start_time = time.time()
    log.info("pipeline.start — file=%s role=%s", resume.filename, target_role)

    # ── Pre-flight ────────────────────────────────────────────────────
    if not target_role or not target_role.strip():
        raise HTTPException(status_code=400, detail="target_role is required.")

    fname = resume.filename or "resume.pdf"
    ext = Path(fname).suffix.lower()
    if ext not in {".pdf", ".docx", ".doc", ".txt", ""}:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type '{ext}'. Accepted: .pdf, .docx, .txt",
        )

    file_bytes = await resume.read()
    if not file_bytes:
        raise HTTPException(status_code=400, detail="Uploaded file is empty.")
    if len(file_bytes) > 15 * 1024 * 1024:
        raise HTTPException(status_code=413, detail="File too large. Max 15 MB.")

    # ── Layer 1: Text Extraction ──────────────────────────────────────
    try:
        raw_text = extract_text(file_bytes, fname)
        log.info("layer1.done — chars=%d", len(raw_text))
    except RuntimeError as e:
        raise HTTPException(status_code=422, detail=str(e))

    if len(raw_text.strip()) < 200:
        raise HTTPException(
            status_code=422,
            detail=(
                f"Resume parsing failed — insufficient content. "
                f"Extracted only {len(raw_text.strip())} characters. "
                f"Please upload a readable PDF or DOCX (not image-only/scanned)."
            ),
        )

    # ── Layer 2: Skill Extraction ─────────────────────────────────────
    skill_result = extract_skills(raw_text)
    user_skills = skill_result["skills"]
    log.info("layer2.done — skills=%d", len(user_skills))

    if not user_skills:
        raise HTTPException(
            status_code=422,
            detail=(
                "No identifiable technical skills found in resume. "
                "Please ensure your resume explicitly lists technologies, "
                "programming languages, and tools you have worked with."
            ),
        )

    # Supplement / override years_of_experience if not provided by form
    years_exp = years_of_experience if years_of_experience is not None else extract_years_experience(raw_text)
    contact   = extract_contact(raw_text)

    # ── Layer 3: Structured Context ───────────────────────────────────
    structured_ctx = build_structured_context(user_skills, target_role.strip(), years_exp)
    log.info("layer3.done — context_len=%d", len(structured_ctx))

    # ── Layer 4: Market Intelligence ─────────────────────────────────
    _role_key, role_data = match_role(target_role.strip())
    if role_data:
        gap_result = compute_skill_gap(user_skills, role_data)
        log.info("layer4.done — alignment=%d%% missing=%d", gap_result["role_alignment_score"], len(gap_result["missing_skills"]))
    else:
        # Unknown role — compute partial alignment based on skill density
        log.warning("layer4.role_not_found — role=%s", target_role)
        gap_result = {
            "critical_missing": [],
            "moderate_missing": [],
            "nice_to_have_missing": [],
            "missing_skills": [],
            "covered_count": len(user_skills),
            "total_required": len(user_skills),
            "role_alignment_score": 50,  # neutral when role unknown
        }

    # ── Layer 5: Confidence Scoring ───────────────────────────────────
    ext_method = ext.lstrip(".") or "pdf"
    confidence_result = compute_confidence(
        text_length=len(raw_text),
        skills_found=len(user_skills),
        extraction_method=ext_method,
        years_exp=years_exp,
        has_contact=bool(contact),
        role_matched=role_data is not None,
        alignment_score=gap_result["role_alignment_score"],
    )
    log.info(
        "layer5.done — confidence=%d grade=%s",
        confidence_result["confidence_score"], confidence_result["grade"],
    )

    # ── Layer 6: Arbitration + Final Report ──────────────────────────
    report = build_final_report(
        user_skills=user_skills,
        skills_by_category=skill_result["by_category"],
        gap_result=gap_result,
        confidence_result=confidence_result,
        target_role=target_role.strip(),
        role_data=role_data,
        years_exp=years_exp,
        contact=contact,
        text_length=len(raw_text),
        filename=fname,
    )

    elapsed = round(time.time() - start_time, 2)
    report["processing_time_seconds"] = elapsed

    log.info(
        "pipeline.complete — skills=%d missing=%d alignment=%d confidence=%d time=%.2fs",
        len(user_skills), len(report["missing_skills"]),
        report["role_alignment_score"], report["confidence_score"], elapsed,
    )

    return JSONResponse(content=report)


# ═══════════════════════════════════════════════════════════════════
# SKILL-GAP COMPATIBILITY ENDPOINT
# (mirrors the path the Node.js proxy calls)
# ═══════════════════════════════════════════════════════════════════

@app.post("/api/skill-gap/analyze-local", tags=["Skill Gap"])
async def skill_gap_analyze_local(
    resume: UploadFile = File(...),
    target_role: str = Form(...),
    years_of_experience: Optional[float] = Form(None),
):
    """Thin wrapper — reuses the main pipeline endpoint."""
    return await analyze_pipeline(
        resume=resume,
        target_role=target_role,
        years_of_experience=years_of_experience,
    )


@app.get("/api/roles", tags=["Reference"])
async def list_roles():
    """Return all roles in the market intelligence database."""
    return {
        "roles": [v["display"] for v in ROLE_DB.values()],
        "total": len(ROLE_DB),
    }


@app.get("/api/skills/ontology", tags=["Reference"])
async def skill_ontology():
    """Return the full skill ontology by category."""
    return {
        "categories": {cat: terms for cat, terms in SKILL_ONTOLOGY.items()},
        "total_skills": sum(len(v) for v in SKILL_ONTOLOGY.values()),
    }


# ══════════════════════════════════════════════════════════════════════
# ENTRY POINT
# ══════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    import uvicorn
    port = int(os.environ.get("PORT", 8001))
    log.info("Starting SkillForge Local AI Engine on port %d", port)
    uvicorn.run("standalone:app", host="0.0.0.0", port=port, reload=True)
